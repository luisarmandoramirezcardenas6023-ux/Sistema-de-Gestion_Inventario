import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import json
import os
import sys
import datetime
import shutil
import openpyxl
from openpyxl.styles import Font
from docx import Document
from fpdf import FPDF

# --- CONFIGURACIÓN DE COLORES ---
COLOR_FONDO_SIDEBAR = "#2c3e50"
COLOR_FONDO_MAIN = "#ecf0f1"
COLOR_TEXTO_OSCURO = "#2c3e50"

# Colores de Botones
COLOR_BTN_NORMAL = "#34495e"
COLOR_BTN_NUEVO = "#f39c12"
COLOR_BTN_MODIFICAR = "#2980b9"
COLOR_BTN_GUARDAR = "#27ae60"
COLOR_BTN_ROJO = "#c0392b"
COLOR_BTN_SALIR = "#1a252f"
COLOR_BTN_REFRESH = "#16a085"
COLOR_BTN_EXCEL = "#217346"
COLOR_BTN_WORD = "#2b5797"
COLOR_BTN_PDF = "#e74c3c"
COLOR_BTN_PRESTAMO = "#8e44ad"
COLOR_BTN_REPORTES_MENU = "#21618C"
COLOR_BTN_VISOR = "#1abc9c"
COLOR_BTN_AYUDA = "#7f8c8d"
COLOR_BTN_RESPALDO = "#d35400"
COLOR_BTN_LEER_RESPALDO = "#17a2b8"
COLOR_BTN_EMPLEADOS = "#8e44ad"

# Color para resaltar selección en reportes
COLOR_HIGHLIGHT_BOX = "#D1F2EB"

# Colores Ordenamiento
COLOR_SORT_OFF = "#bdc3c7"
COLOR_SORT_ON = "#2c3e50"
COLOR_SORT_FG_OFF = "#7f8c8d"
COLOR_SORT_FG_ON = "#ffffff"

# --- LÓGICA DE DATOS Y RUTAS ---
if getattr(sys, 'frozen', False):
    CARPETA_ACTUAL = os.path.dirname(sys.executable)
else:
    CARPETA_ACTUAL = os.path.dirname(os.path.abspath(__file__))

ARCHIVO_DATOS = os.path.join(CARPETA_ACTUAL, "inventario_taller.json")
ARCHIVO_LOG = os.path.join(CARPETA_ACTUAL, "historial_global.json")
CARPETA_RESPALDOS = os.path.join(CARPETA_ACTUAL, "Respaldos")


# --- CLASE PDF PERSONALIZADA ---
class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, 'SISTEMA DE GESTION DE INVENTARIO', 0, 1, 'C')
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Pagina {self.page_no()}', 0, 0, 'C')

    def clean_text(self, text):
        if not text: return ""
        text = str(text)
        replacements = {
            "📤": ">>", "📥": "<<", "✅": "OK", "✨": "+",
            "🗑️": "DEL", "✏️": "EDT", "💾": "SAV", "🔄": "MOV",
            "🛡️": "", "📊": "", "📅": "", "👁️": "", "👷": "", "❌": "NO", "⚠️": "WARN"
        }
        for k, v in replacements.items():
            text = text.replace(k, v)
        return text.encode('latin-1', 'replace').decode('latin-1')


# --- FUNCIONES DE CARGA DE DATOS ---
def cargar_datos():
    if not os.path.exists(ARCHIVO_DATOS):
        try:
            with open(ARCHIVO_DATOS, "w", encoding="utf-8") as archivo:
                json.dump({}, archivo, indent=4)
            return {}
        except:
            return {}
    else:
        try:
            with open(ARCHIVO_DATOS, "r", encoding="utf-8") as archivo:
                return json.load(archivo)
        except:
            return {}


def guardar_datos(inventario):
    try:
        with open(ARCHIVO_DATOS, "w", encoding="utf-8") as archivo:
            json.dump(inventario, archivo, indent=4)
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo guardar inventario: {e}")


# --- FUNCIONES DE LOG GLOBAL ---
def cargar_log_global():
    if not os.path.exists(ARCHIVO_LOG): return []
    try:
        with open(ARCHIVO_LOG, "r", encoding="utf-8") as f:
            return json.load(f)
    except:
        return []


def registrar_accion_global(accion, codigo, nombre, detalle_extra=""):
    log = cargar_log_global()
    ahora = datetime.datetime.now()
    nuevo_evento = {
        "fecha": ahora.strftime("%d/%m/%Y"),
        "hora": ahora.strftime("%H:%M:%S"),
        "accion": accion,
        "codigo": codigo,
        "nombre": nombre,
        "detalle": detalle_extra
    }
    log.append(nuevo_evento)
    try:
        with open(ARCHIVO_LOG, "w", encoding="utf-8") as f:
            json.dump(log, f, indent=4)
    except:
        pass


# --- CLASE PRINCIPAL ---
class SistemaInventario:
    def __init__(self, root):
        self.root = root
        self.root.title("Sistema de Inventario (Almacen)")
        self.root.geometry("1280x800")
        self.root.configure(bg=COLOR_FONDO_MAIN)

        self.pixel = tk.PhotoImage(width=1, height=1)

        self.inventario = cargar_datos()
        self.orden_actual = "id"
        self.modo_actual = "lectura"

        self.ejecutar_respaldo_inicio()

        self.configurar_estilos()
        self.construir_interfaz()

        self.actualizar_estilo_botones_sort()
        self.bloquear_campos()
        self.actualizar_botones_sidebar()
        self.limpiar_campos_visual()

    def configurar_estilos(self):
        style = ttk.Style()
        style.theme_use("clam")
        style.configure("Treeview.Heading", font=("Segoe UI", 10, "bold"),
                        background=COLOR_FONDO_SIDEBAR, foreground="white", padding=5)
        style.configure("Treeview", font=("Segoe UI", 9), rowheight=25)
        style.map("Treeview", background=[("selected", "#3498db")])
        self.root.option_add('*Entry.disabledBackground', '#ecf0f1')
        self.root.option_add('*Entry.disabledForeground', '#7f8c8d')

    def ejecutar_respaldo_inicio(self):
        if not os.path.exists(ARCHIVO_DATOS): return
        if not self.inventario: return
        if not os.path.exists(CARPETA_RESPALDOS):
            try:
                os.mkdir(CARPETA_RESPALDOS)
            except:
                return

        timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        nombre_bak = f"Respaldo_Auto_{timestamp}.json"
        ruta_destino = os.path.join(CARPETA_RESPALDOS, nombre_bak)
        try:
            shutil.copy2(ARCHIVO_DATOS, ruta_destino)
        except:
            pass

    def crear_respaldo_manual(self):
        if not self.inventario: return messagebox.showwarning("Vacío", "No hay datos para respaldar.")

        if not messagebox.askyesno("Confirmar Respaldo", "¿Estás seguro de crear una copia de seguridad ahora?"):
            return

        if not os.path.exists(CARPETA_RESPALDOS): os.mkdir(CARPETA_RESPALDOS)
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        nombre_bak = f"Respaldo_MANUAL_{timestamp}.json"
        ruta_destino = os.path.join(CARPETA_RESPALDOS, nombre_bak)
        try:
            shutil.copy2(ARCHIVO_DATOS, ruta_destino)
            messagebox.showinfo("Respaldo Exitoso",
                                f"✅ Copia creada correctamente en:\nCarpeta 'Respaldos'\nArchivo: {nombre_bak}")
        except Exception as e:
            messagebox.showerror("Error", f"Fallo al crear respaldo: {e}")

    def abrir_visor_respaldos(self):
        if not messagebox.askyesno("Confirmar Visor", "¿Deseas buscar y leer un respaldo anterior?"):
            return

        if not os.path.exists(CARPETA_RESPALDOS):
            return messagebox.showinfo("Aviso", "Aún no hay carpeta de respaldos.")

        ruta = filedialog.askopenfilename(
            initialdir=CARPETA_RESPALDOS,
            title="Seleccionar Respaldo para Visualizar",
            filetypes=[("Archivos de Respaldo", "*.json")]
        )

        if not ruta: return

        try:
            with open(ruta, 'r', encoding='utf-8') as f:
                datos_backup = json.load(f)
        except Exception as e:
            return messagebox.showerror("Error", f"No se pudo leer el archivo: {e}")

        ventana = tk.Toplevel(self.root)
        ventana.title(f"VISOR HISTÓRICO - {os.path.basename(ruta)}")
        ventana.geometry("1100x600")
        ventana.configure(bg="#f1c40f")

        lbl_info = tk.Label(ventana, text=f"⚠️ MODO LECTURA: Visualizando archivo histórico: {os.path.basename(ruta)}",
                            bg="#f1c40f", fg="black", font=("Segoe UI", 12, "bold"), pady=10)
        lbl_info.pack(fill=tk.X)

        frame_tabla = tk.Frame(ventana, bg="white")
        frame_tabla.pack(fill=tk.BOTH, expand=True, padx=15, pady=(0, 15))

        scroll_y = tk.Scrollbar(frame_tabla, orient=tk.VERTICAL)
        cols = ("ID", "Codigo", "Nombre", "Cantidad", "Gabinete", "UltimoMov", "Descripcion")
        tree = ttk.Treeview(frame_tabla, columns=cols, show='headings', yscrollcommand=scroll_y.set)

        scroll_y.config(command=tree.yview)
        scroll_y.pack(side=tk.RIGHT, fill=tk.Y)
        tree.pack(fill=tk.BOTH, expand=True)

        tree.heading("ID", text="ID");
        tree.column("ID", width=40, anchor=tk.CENTER)
        tree.heading("Codigo", text="CÓD.");
        tree.column("Codigo", width=80, anchor=tk.CENTER)
        tree.heading("Nombre", text="NOMBRE PIEZA");
        tree.column("Nombre", width=200)
        tree.heading("Cantidad", text="CANT.");
        tree.column("Cantidad", width=60, anchor=tk.CENTER)
        tree.heading("Gabinete", text="GAB.");
        tree.column("Gabinete", width=80, anchor=tk.CENTER)
        tree.heading("UltimoMov", text="ÚLTIMO MOV. REGISTRADO");
        tree.column("UltimoMov", width=200)
        tree.heading("Descripcion", text="DESCRIPCIÓN");
        tree.column("Descripcion", width=250)

        items_ordenados = sorted(datos_backup.items(), key=lambda x: int(x[0]) if x[0].isdigit() else 0)

        for pid, d in items_ordenados:
            historial = d.get('historial', [])
            ultimo = historial[0] if historial else "Sin movimientos"
            tree.insert("", tk.END, values=(pid, d.get('codigo', ''), d['nombre'], d['cantidad'], d['gabinete'], ultimo,
                                            d['descripcion']))

        btn_cerrar = tk.Button(ventana, text="Cerrar Visor", command=ventana.destroy,
                               bg="#34495e", fg="white", font=("Segoe UI", 10, "bold"), height=2)
        btn_cerrar.pack(fill=tk.X, padx=15, pady=(0, 15))

    def construir_interfaz(self):
        self.sidebar = tk.Frame(self.root, bg=COLOR_FONDO_SIDEBAR, width=320)
        self.sidebar.pack(side=tk.LEFT, fill=tk.Y)
        self.sidebar.pack_propagate(False)

        content_sidebar = tk.Frame(self.sidebar, bg=COLOR_FONDO_SIDEBAR)
        content_sidebar.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

        frame_top = tk.Frame(content_sidebar, bg=COLOR_FONDO_SIDEBAR)
        frame_top.pack(fill=tk.X, padx=10, pady=(5, 0))
        self.btn_ayuda = tk.Button(frame_top, text="❓ AYUDA / MANUAL", command=self.abrir_ayuda,
                                   bg=COLOR_BTN_AYUDA, fg="white", font=("Segoe UI", 7, "bold"), bd=0, cursor="hand2")
        self.btn_ayuda.pack(side=tk.RIGHT)

        lbl_brand = tk.Label(content_sidebar, text="INVENTARIO", font=("Segoe UI", 18, "bold"),
                             bg=COLOR_FONDO_SIDEBAR, fg="white")
        lbl_brand.pack(pady=(5, 5))
        tk.Frame(content_sidebar, bg="white", height=1, width=280).pack(pady=5)

        frame_inputs = tk.Frame(content_sidebar, bg=COLOR_FONDO_SIDEBAR)
        frame_inputs.pack(fill=tk.X, pady=5)

        self.crear_label_input(frame_inputs, "ID Sistema:")
        self.entry_id = self.crear_entry(frame_inputs, readonly=True)
        self.crear_label_input(frame_inputs, "Código:")
        self.entry_codigo = self.crear_entry(frame_inputs, solo_numeros=True)
        self.crear_label_input(frame_inputs, "Pieza:")
        self.entry_nombre = self.crear_entry(frame_inputs)
        self.crear_label_input(frame_inputs, "Stock:")
        self.entry_cantidad = self.crear_entry(frame_inputs, solo_numeros=True)
        self.crear_label_input(frame_inputs, "Gabinete:")
        self.entry_gabinete = self.crear_entry(frame_inputs)
        self.crear_label_input(frame_inputs, "Descripción:")
        self.entry_desc = self.crear_entry(frame_inputs)

        self.campos_editables = [self.entry_codigo, self.entry_nombre, self.entry_cantidad, self.entry_gabinete,
                                 self.entry_desc]

        frame_botones = tk.Frame(content_sidebar, bg=COLOR_FONDO_SIDEBAR)
        frame_botones.pack(fill=tk.X, padx=15, pady=10)

        self.btn_nuevo = self.crear_boton(frame_botones, "✨ NUEVO / LIMPIAR", self.accion_boton_nuevo, COLOR_BTN_NORMAL)
        self.btn_modificar = self.crear_boton(frame_botones, "✏️ MODIFICAR", self.accion_boton_modificar,
                                              COLOR_BTN_NORMAL)
        self.btn_eliminar = self.crear_boton(frame_botones, "🗑️ ELIMINAR", self.eliminar_pieza, COLOR_BTN_ROJO)

        tk.Frame(frame_botones, bg=COLOR_FONDO_SIDEBAR, height=2).pack()
        self.btn_prestamo = self.crear_boton(frame_botones, "🔄 PRÉSTAMO / DEVOLUCIÓN", self.abrir_ventana_prestamos,
                                             COLOR_BTN_PRESTAMO)
        tk.Frame(frame_botones, bg=COLOR_FONDO_SIDEBAR, height=2).pack()
        self.btn_reportes = self.crear_boton(frame_botones, "📂 CENTRO DE REPORTES", self.abrir_centro_reportes,
                                             COLOR_BTN_REPORTES_MENU)
        tk.Frame(frame_botones, bg=COLOR_FONDO_SIDEBAR, height=20).pack()
        self.btn_salir = self.crear_boton(frame_botones, "🚪 SALIR DEL SISTEMA", self.salir_sistema, COLOR_BTN_SALIR)

        self.main_area = tk.Frame(self.root, bg=COLOR_FONDO_MAIN)
        self.main_area.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=15, pady=15)

        header_frame = tk.Frame(self.main_area, bg=COLOR_FONDO_MAIN)
        header_frame.pack(fill=tk.X, pady=(0, 5))

        lbl_lista = tk.Label(header_frame, text="Monitor de Inventario", font=("Segoe UI", 16, "bold"),
                             bg=COLOR_FONDO_MAIN, fg=COLOR_TEXTO_OSCURO)
        lbl_lista.pack(side=tk.LEFT)

        frame_search = tk.Frame(header_frame, bg=COLOR_FONDO_MAIN)
        frame_search.pack(side=tk.LEFT, padx=20)

        tk.Label(frame_search, text="Buscar (Cód/Nombre):", bg=COLOR_FONDO_MAIN, fg="#7f8c8d",
                 font=("Segoe UI", 9, "bold")).pack(side=tk.LEFT)
        self.entry_buscar = tk.Entry(frame_search, font=("Segoe UI", 9), width=20)
        self.entry_buscar.pack(side=tk.LEFT, padx=5)
        self.entry_buscar.bind('<Return>', self.realizar_busqueda)

        self.btn_search = tk.Button(frame_search, text="🔍", command=self.realizar_busqueda,
                                    bg=COLOR_BTN_MODIFICAR, fg="white", relief="flat", cursor="hand2", width=3)
        self.btn_search.pack(side=tk.LEFT)

        frame_sort = tk.Frame(header_frame, bg=COLOR_FONDO_MAIN)
        frame_sort.pack(side=tk.RIGHT)

        tk.Label(frame_sort, text="Ordenar:", bg=COLOR_FONDO_MAIN, fg="#7f8c8d", font=("Segoe UI", 8)).pack(
            side=tk.LEFT, padx=5)
        self.btn_sort_id = self.crear_boton_sort(frame_sort, "ID", "id")
        self.btn_sort_nombre = self.crear_boton_sort(frame_sort, "A-Z", "nombre")
        self.btn_sort_cantidad = self.crear_boton_sort(frame_sort, "1-9", "cantidad")

        # --- NUEVO BOTÓN: ORDENAR POR RECIENTE ---
        self.btn_sort_reciente = self.crear_boton_sort(frame_sort, "🕒 Reciente", "reciente")

        tk.Frame(frame_sort, bg=COLOR_FONDO_MAIN, width=10).pack(side=tk.LEFT)

        self.btn_ref = tk.Button(frame_sort, text="🔄 Refrescar Tabla", command=self.accion_refrescar_manual,
                                 bg=COLOR_BTN_REFRESH, fg="white", font=("Segoe UI", 8, "bold"),
                                 relief="flat", bd=0, cursor="hand2", padx=8, pady=3)
        self.btn_ref.pack(side=tk.LEFT, padx=2)

        frame_tabla = tk.Frame(self.main_area, bg="white", bd=1, relief="solid")
        frame_tabla.pack(fill=tk.BOTH, expand=True)

        scroll_y = tk.Scrollbar(frame_tabla, orient=tk.VERTICAL)
        scroll_x = tk.Scrollbar(frame_tabla, orient=tk.HORIZONTAL)

        columnas = ("ID", "Codigo", "Nombre", "Cantidad", "Gabinete", "Estatus", "Descripcion")

        self.tabla = ttk.Treeview(frame_tabla, columns=columnas, show='headings',
                                  yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)

        scroll_y.config(command=self.tabla.yview);
        scroll_y.pack(side=tk.RIGHT, fill=tk.Y)
        scroll_x.config(command=self.tabla.xview);
        scroll_x.pack(side=tk.BOTTOM, fill=tk.X)

        self.tabla.heading("ID", text="ID")
        self.tabla.heading("Codigo", text="CÓD.")
        self.tabla.heading("Nombre", text="NOMBRE PIEZA")
        self.tabla.heading("Cantidad", text="CANT.")
        self.tabla.heading("Gabinete", text="GAB.")
        self.tabla.heading("Estatus", text="ESTATUS HOY")
        self.tabla.heading("Descripcion", text="DESCRIPCIÓN")

        self.tabla.column("ID", width=35, anchor=tk.CENTER)
        self.tabla.column("Codigo", width=70, anchor=tk.CENTER)
        self.tabla.column("Nombre", width=180)
        self.tabla.column("Cantidad", width=50, anchor=tk.CENTER)
        self.tabla.column("Gabinete", width=60, anchor=tk.CENTER)
        self.tabla.column("Estatus", width=140, anchor=tk.CENTER)
        self.tabla.column("Descripcion", width=200)

        self.tabla.pack(fill=tk.BOTH, expand=True)
        self.tabla.bind("<ButtonRelease-1>", self.seleccionar_item)

        lbl_detalles = tk.Label(self.main_area, text="Ficha Técnica e Historial Completo",
                                font=("Segoe UI", 12, "bold"), bg=COLOR_FONDO_MAIN, fg=COLOR_TEXTO_OSCURO)
        lbl_detalles.pack(anchor="w", pady=(10, 5))

        frame_detalles = tk.Frame(self.main_area, bg="white", bd=1, relief="solid", height=150)
        frame_detalles.pack(fill=tk.X)
        frame_detalles.pack_propagate(False)

        self.txt_detalles = tk.Text(frame_detalles, font=("Consolas", 9), bg="white", relief="flat", padx=10, pady=10)
        self.txt_detalles.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        self.txt_detalles.tag_config("azul", foreground="blue")
        self.txt_detalles.tag_config("verde", foreground="green")
        self.txt_detalles.tag_config("negro", foreground="black")

        scroll_det = tk.Scrollbar(frame_detalles, command=self.txt_detalles.yview)
        scroll_det.pack(side=tk.RIGHT, fill=tk.Y)
        self.txt_detalles.config(yscrollcommand=scroll_det.set, state=tk.DISABLED)

        self.refrescar_tabla()

    def crear_boton(self, parent, texto, comando, color):
        btn = tk.Button(parent, text=texto, command=comando, bg=color, fg="white",
                        font=("Segoe UI", 8, "bold"), relief="flat", bd=0, cursor="hand2",
                        image=self.pixel, compound="center", height=30, width=280)
        btn.pack(fill=tk.X, pady=2)
        return btn

    def crear_label_input(self, parent, texto):
        tk.Label(parent, text=texto, font=("Segoe UI", 8, "bold"), bg=COLOR_FONDO_SIDEBAR, fg="#ecf0f1").pack(
            anchor="w", padx=15, pady=(2, 0))

    def crear_entry(self, parent, readonly=False, solo_numeros=False):
        entry = ttk.Entry(parent, font=("Segoe UI", 9))
        if solo_numeros: entry.config(validate='key',
                                      validatecommand=(self.root.register(self.validar_solo_numeros), '%P'))
        entry.pack(fill=tk.X, padx=15, pady=2)
        if readonly: entry.configure(state="readonly")
        return entry

    def crear_boton_sort(self, parent, texto, tipo_orden):
        btn = tk.Button(parent, text=texto, command=lambda: self.cambiar_orden(tipo_orden), bg=COLOR_SORT_OFF,
                        fg=COLOR_SORT_FG_OFF, font=("Segoe UI", 8, "bold"), relief="flat", bd=0, cursor="hand2", padx=8,
                        pady=1)
        btn.pack(side=tk.LEFT, padx=2)
        btn.bind("<Enter>", lambda e: btn.config(bg="#95a5a6", fg="white"));
        btn.bind("<Leave>",
                 lambda e: btn.config(bg=COLOR_SORT_ON if self.orden_actual == tipo_orden else COLOR_SORT_OFF,
                                      fg=COLOR_SORT_FG_ON if self.orden_actual == tipo_orden else COLOR_SORT_FG_OFF))
        return btn

    def validar_solo_numeros(self, texto):
        return texto.isdigit() or texto == ""

    def verificar_bloqueo(self):
        if self.modo_actual == "editar":
            messagebox.showwarning("Modo Edición",
                                   "⚠️ Termina de guardar o cancelar los cambios antes de hacer otra cosa.")
            return True
        return False

    def actualizar_botones_sidebar(self):
        if self.modo_actual == "lectura":
            self.btn_nuevo.config(text="✨ NUEVO / LIMPIAR", bg=COLOR_BTN_NUEVO, command=self.accion_boton_nuevo)
            self.btn_modificar.config(text="✏️ MODIFICAR", bg=COLOR_BTN_MODIFICAR)
        elif self.modo_actual == "nuevo":
            self.btn_nuevo.config(text="💾 GUARDAR NUEVO", bg=COLOR_BTN_GUARDAR, command=self.agregar_pieza)
        elif self.modo_actual == "editar":
            self.btn_nuevo.config(text="❌ CANCELAR EDICIÓN", bg=COLOR_BTN_ROJO, command=self.cancelar_edicion)
            self.btn_modificar.config(text="💾 GUARDAR CAMBIOS", bg=COLOR_BTN_GUARDAR)

    def actualizar_estilo_botones_sort(self):
        # Actualizamos TODOS los botones, incluido el nuevo de RECIENTE
        for btn in [self.btn_sort_id, self.btn_sort_nombre, self.btn_sort_cantidad, self.btn_sort_reciente]:
            btn.config(bg=COLOR_SORT_OFF, fg=COLOR_SORT_FG_OFF)

        if self.orden_actual == "id":
            self.btn_sort_id.config(bg=COLOR_SORT_ON, fg=COLOR_SORT_FG_ON)
        elif self.orden_actual == "nombre":
            self.btn_sort_nombre.config(bg=COLOR_SORT_ON, fg=COLOR_SORT_FG_ON)
        elif self.orden_actual == "cantidad":
            self.btn_sort_cantidad.config(bg=COLOR_SORT_ON, fg=COLOR_SORT_FG_ON)
        elif self.orden_actual == "reciente":
            self.btn_sort_reciente.config(bg=COLOR_SORT_ON, fg=COLOR_SORT_FG_ON)

    def bloquear_campos(self):
        for campo in self.campos_editables: campo.config(state="readonly")

    def desbloquear_campos(self):
        for campo in self.campos_editables: campo.config(state="normal")

    def limpiar_campos_logica(self):
        self.entry_id.config(state='normal')
        self.entry_id.delete(0, tk.END)
        self.entry_id.insert(0, self.generar_proximo_id())
        self.entry_id.config(state='readonly')
        self.desbloquear_campos()
        for c in self.campos_editables: c.delete(0, tk.END)
        self.txt_detalles.config(state=tk.NORMAL)
        self.txt_detalles.delete(1.0, tk.END)
        self.txt_detalles.insert(tk.END, "Modo Creación: Ingresa los datos numéricos.")
        self.txt_detalles.config(state=tk.DISABLED)

    def limpiar_campos_visual(self):
        for c in self.campos_editables:
            c.config(state="normal")
            c.delete(0, tk.END)
            c.config(state="readonly")
        self.txt_detalles.config(state=tk.NORMAL)
        self.txt_detalles.delete(1.0, tk.END)
        self.txt_detalles.insert(tk.END, "Bienvenido. Selecciona una herramienta o crea una nueva.")
        self.txt_detalles.config(state=tk.DISABLED)

    def validar_codigo_duplicado(self, codigo, id_actual=None):
        codigo = str(codigo).strip().lower()
        if not codigo: return False
        for pid, data in self.inventario.items():
            if id_actual and str(pid) == str(id_actual): continue
            codigo_existente = str(data.get('codigo', '')).strip().lower()
            if codigo_existente == codigo:
                return data['nombre']
        return None

    def accion_boton_nuevo(self):
        if self.verificar_bloqueo(): return
        self.refrescar_tabla()
        self.modo_actual = "nuevo"
        self.limpiar_campos_logica()
        if self.tabla.selection(): self.tabla.selection_remove(self.tabla.selection())
        self.actualizar_botones_sidebar()

    def cancelar_edicion(self):
        if messagebox.askyesno("Cancelar", "¿Descartar cambios y salir de la edición?"):
            self.modo_actual = "lectura"
            self.bloquear_campos()
            self.actualizar_botones_sidebar()
            self.seleccionar_item(None)

    def accion_boton_modificar(self):
        if self.modo_actual == "nuevo": return
        if not self.entry_id.get(): return messagebox.showwarning("Aviso", "Selecciona una pieza.")
        if self.modo_actual == "lectura":
            if messagebox.askyesno("Habilitar", "¿Modificar esta pieza?"):
                self.modo_actual = "editar"
                self.desbloquear_campos()
                self.entry_nombre.focus()
                self.actualizar_botones_sidebar()
        elif self.modo_actual == "editar":
            self.actualizar_pieza()

    def generar_proximo_id(self):
        ids = [int(k) for k in self.inventario if k.isdigit()]
        return str(max(ids) + 1) if ids else "1"

    def accion_refrescar_manual(self):
        if self.verificar_bloqueo(): return
        self.refrescar_tabla(self.obtener_id_actual_seleccionado())

    def obtener_id_actual_seleccionado(self):
        sel = self.tabla.selection()
        return str(self.tabla.item(sel[0])['values'][0]) if sel else None

    def realizar_busqueda(self, event=None):
        if self.verificar_bloqueo(): return
        termino = self.entry_buscar.get().strip().lower()
        if not termino: return messagebox.showwarning("Buscador", "Ingresa un código o nombre.")
        resultados = []
        for pid, data in self.inventario.items():
            codigo = str(data.get('codigo', '')).lower()
            nombre = str(data['nombre']).lower()
            if termino in codigo or termino in nombre:
                resultados.append((pid, data))

        if len(resultados) == 0:
            messagebox.showerror("Sin resultados", f"No se encontró nada con '{termino}'.")
        elif len(resultados) == 1:
            pid, data = resultados[0]
            self.refrescar_tabla(id_seleccionado=pid)
            messagebox.showinfo("Encontrado", f"✅ Pieza: {data['nombre']}")
            self.entry_buscar.delete(0, tk.END)
        else:
            self.mostrar_resultados_busqueda(resultados)

    def mostrar_resultados_busqueda(self, resultados):
        ventana = tk.Toplevel(self.root)
        ventana.title(f"{len(resultados)} Coincidencias")
        ventana.geometry("400x400")
        ventana.configure(bg=COLOR_FONDO_MAIN)
        ventana.grab_set()
        tk.Label(ventana, text="Selecciona la pieza que buscas:", font=("Segoe UI", 10, "bold"),
                 bg=COLOR_FONDO_MAIN).pack(pady=10)
        listbox = tk.Listbox(ventana, font=("Segoe UI", 10), width=50, height=15)
        listbox.pack(padx=20, pady=5)
        mapa_ids = []
        for pid, data in resultados:
            texto = f"[{data.get('codigo', '')}] {data['nombre']} (Stock: {data['cantidad']})"
            listbox.insert(tk.END, texto)
            mapa_ids.append(pid)

        def seleccionar():
            idx = listbox.curselection()
            if not idx: return
            pid_elegido = mapa_ids[idx[0]]
            self.refrescar_tabla(id_seleccionado=pid_elegido)
            self.entry_buscar.delete(0, tk.END)
            ventana.destroy()

        tk.Button(ventana, text="Ver Pieza", command=seleccionar, bg=COLOR_BTN_MODIFICAR, fg="white",
                  font=("Segoe UI", 10, "bold")).pack(pady=10)
        listbox.bind('<Double-Button-1>', lambda e: seleccionar())

    def refrescar_tabla(self, id_seleccionado=None):
        for item in self.tabla.get_children(): self.tabla.delete(item)
        items = self.inventario.items()

        # --- LÓGICA DE ORDENAMIENTO ---
        if self.orden_actual == "nombre":
            items_ordenados = sorted(items, key=lambda x: x[1]['nombre'].lower())
        elif self.orden_actual == "cantidad":
            items_ordenados = sorted(items, key=lambda x: int(x[1]['cantidad']))
        elif self.orden_actual == "reciente":
            # Función auxiliar para obtener la fecha del último movimiento
            def obtener_fecha_ultimo_mov(item):
                hist = item[1].get('historial', [])
                if not hist: return datetime.datetime.min  # Si no hay historial, al fondo
                try:
                    # Formato esperado: "📤 07/01/2026 15:30:00 | ..."
                    # Partimos por el pipe "|" y luego quitamos el icono
                    partes = hist[0].split("|")[0].strip().split(" ", 1)
                    if len(partes) > 1:
                        fecha_str = partes[1]  # "07/01/2026 15:30:00"
                        return datetime.datetime.strptime(fecha_str, "%d/%m/%Y %H:%M:%S")
                    return datetime.datetime.min
                except:
                    return datetime.datetime.min

            # Ordenamos DESCENDENTE (lo más nuevo arriba)
            items_ordenados = sorted(items, key=obtener_fecha_ultimo_mov, reverse=True)
        else:
            items_ordenados = sorted(items, key=lambda x: int(x[0]) if x[0].isdigit() else 0)

        # --- SEMÁFORO DE INVENTARIO (CONFIGURACIÓN) ---
        self.tabla.tag_configure('stock_critico', background='#e74c3c', foreground='white')  # Rojo
        self.tabla.tag_configure('stock_bajo', background='#f1c40f', foreground='black')  # Amarillo

        for pid, d in items_ordenados:
            estatus_hoy = self.obtener_estatus_hoy_texto(d.get('historial', []))

            # --- SEMÁFORO DE INVENTARIO (LÓGICA) ---
            cantidad = int(d['cantidad'])
            tag_fila = ()  # Sin color por defecto

            if cantidad <= 2:
                tag_fila = ('stock_critico',)  # Menos de 2: ROJO
            elif cantidad <= 5:
                tag_fila = ('stock_bajo',)  # Entre 3 y 5: AMARILLO

            item_id = self.tabla.insert("", tk.END,
                                        values=(pid, d.get('codigo', ''), d['nombre'], d['cantidad'], d['gabinete'],
                                                estatus_hoy, d['descripcion']),
                                        tags=tag_fila)  # Agregar el tag aquí

            if id_seleccionado and str(pid) == str(id_seleccionado):
                self.tabla.selection_set(item_id);
                self.tabla.see(item_id);
                self.seleccionar_item(None)

    def seleccionar_item(self, event):
        sel = self.tabla.selection()
        if not sel: return
        if self.modo_actual == "editar": return
        if self.modo_actual == "nuevo":
            self.modo_actual = "lectura"
            self.actualizar_botones_sidebar()

        vals = self.tabla.item(sel[0], 'values')
        self.entry_id.config(state='normal')
        self.entry_id.delete(0, tk.END)
        self.entry_id.insert(0, vals[0])
        self.entry_id.config(state='readonly')
        self.desbloquear_campos()
        self.entry_codigo.delete(0, tk.END);
        self.entry_codigo.insert(0, vals[1])
        self.entry_nombre.delete(0, tk.END);
        self.entry_nombre.insert(0, vals[2])
        self.entry_cantidad.delete(0, tk.END);
        self.entry_cantidad.insert(0, vals[3])
        self.entry_gabinete.delete(0, tk.END);
        self.entry_gabinete.insert(0, vals[4])
        self.entry_desc.delete(0, tk.END);
        self.entry_desc.insert(0, vals[6])
        self.bloquear_campos()

        pid = vals[0]
        historial = self.inventario[pid].get('historial', [])
        self.txt_detalles.config(state=tk.NORMAL);
        self.txt_detalles.delete(1.0, tk.END)
        info_header = f"--- FICHA TÉCNICA ---\nID: {pid} | CÓDIGO: {vals[1]} | PIEZA: {vals[2]} | UBICACIÓN: {vals[4]}\nDESCRIPCIÓN: {vals[6]}\n\n--- HISTORIAL DETALLADO ---\n"
        self.txt_detalles.insert(tk.END, info_header, "negro")

        if not historial:
            self.txt_detalles.insert(tk.END, "Sin movimientos registrados.", "negro")
        else:
            for linea in historial:
                if "SALIDA" in linea:
                    self.txt_detalles.insert(tk.END, linea + "\n", "azul")
                elif "ENTRADA" in linea:
                    self.txt_detalles.insert(tk.END, linea + "\n", "verde")
                else:
                    self.txt_detalles.insert(tk.END, linea + "\n", "negro")
        self.txt_detalles.config(state=tk.DISABLED)

    def agregar_pieza(self):
        pid, cod, nom, cant, gab, desc = self.entry_id.get(), self.entry_codigo.get().strip(), self.entry_nombre.get().strip(), self.entry_cantidad.get().strip(), self.entry_gabinete.get().strip(), self.entry_desc.get().strip()

        if not cod or not nom or not cant or not gab:
            return messagebox.showwarning("Incompleto", "⚠️ Faltan datos obligatorios.")

        nombre_duplicado = self.validar_codigo_duplicado(cod)
        if nombre_duplicado:
            return messagebox.showerror("Código Repetido",
                                        f"⛔ El código '{cod}' ya está registrado.\nPertenece a: {nombre_duplicado}")

        if not messagebox.askyesno("Guardar", f"¿Registrar '{nom}'?"): return

        try:
            self.inventario[pid] = {"codigo": cod, "nombre": nom, "cantidad": int(cant), "gabinete": gab,
                                    "descripcion": desc, "historial": []}

            registrar_accion_global("CREACIÓN", cod, nom, f"Stock Inicial: {cant}")

            guardar_datos(self.inventario);
            self.refrescar_tabla(id_seleccionado=pid);
            self.modo_actual = "lectura"
            self.bloquear_campos()
            self.actualizar_botones_sidebar()
            messagebox.showinfo("Éxito", "Pieza guardada.")
        except ValueError:
            messagebox.showerror("Error", "Cantidad debe ser número")

    def actualizar_pieza(self):
        pid, cod, nom, cant, gab, desc = self.entry_id.get(), self.entry_codigo.get().strip(), self.entry_nombre.get().strip(), self.entry_cantidad.get().strip(), self.entry_gabinete.get().strip(), self.entry_desc.get().strip()

        nombre_duplicado = self.validar_codigo_duplicado(cod, id_actual=pid)
        if nombre_duplicado:
            return messagebox.showerror("Código Repetido",
                                        f"⛔ El código '{cod}' ya está en uso por otra pieza:\n{nombre_duplicado}")

        if not messagebox.askyesno("Actualizar", f"¿Guardar cambios en '{nom}'?"): return
        try:
            historial_previo = self.inventario[pid].get('historial', [])
            self.inventario[pid] = {"codigo": cod, "nombre": nom, "cantidad": int(cant), "gabinete": gab,
                                    "descripcion": desc, "historial": historial_previo}
            guardar_datos(self.inventario);
            self.refrescar_tabla(id_seleccionado=pid);
            self.modo_actual = "lectura"
            self.bloquear_campos()
            self.actualizar_botones_sidebar()
            messagebox.showinfo("Éxito", "Cambios guardados.")
        except ValueError:
            messagebox.showerror("Error", "Cantidad debe ser número")

    def eliminar_pieza(self):
        if self.verificar_bloqueo(): return
        seleccion = self.tabla.selection()
        if not seleccion: return messagebox.showwarning("Aviso",
                                                        "⚠️ Por favor, selecciona una herramienta de la lista para eliminar.")
        pid = self.entry_id.get();
        if not pid: return messagebox.showwarning("Aviso", "No hay un ID seleccionado.")

        if messagebox.askyesno("Eliminar", f"¿Estás seguro de ELIMINAR permanentemente la pieza ID {pid}?"):
            data = self.inventario[pid]
            registrar_accion_global("ELIMINACIÓN", data.get('codigo', ''), data['nombre'], "Pieza dada de baja")
            del self.inventario[pid];
            guardar_datos(self.inventario);
            self.refrescar_tabla();
            self.modo_actual = "lectura"
            self.limpiar_campos_visual()
            self.actualizar_botones_sidebar()

    # --- NUEVA VERSIÓN DE PRÉSTAMOS (CORREGIDA Y COMPACTA) ---
    def abrir_ventana_prestamos(self):
        if self.verificar_bloqueo(): return

        ventana = tk.Toplevel(self.root)
        ventana.title("🛒 Gestión Masiva (Con Buscador Visual)")
        ventana.geometry("800x600")
        ventana.configure(bg=COLOR_FONDO_MAIN)
        ventana.grab_set()

        carrito = []

        # --- SECCIÓN 1: EMPLEADO (Arriba) ---
        frame_emp = tk.Frame(ventana, bg=COLOR_FONDO_MAIN)
        frame_emp.pack(fill=tk.X, padx=20, pady=10)

        tk.Label(frame_emp, text="Empleado:", font=("Segoe UI", 10, "bold"),
                 bg=COLOR_FONDO_MAIN).pack(side=tk.LEFT)

        # --- AQUÍ ESTÁ LA MEJORA: VALIDACIÓN SOLO NÚMEROS ---
        vcmd = (self.root.register(self.validar_solo_numeros), '%P')
        entry_empleado = ttk.Entry(frame_emp, font=("Segoe UI", 11), width=20, justify="center",
                                   validate='key', validatecommand=vcmd)
        entry_empleado.pack(side=tk.LEFT, padx=10)
        entry_empleado.focus()

        tk.Frame(ventana, bg="#bdc3c7", height=2).pack(fill=tk.X, padx=10, pady=5)

        # --- SECCIÓN 2: ÁREA DE TRABAJO DIVIDIDA ---
        frame_middle = tk.Frame(ventana, bg=COLOR_FONDO_MAIN)
        frame_middle.pack(fill=tk.BOTH, expand=True, padx=20, pady=5)

        # -- SUB-PANEL IZQUIERDO: BUSCADOR DE STOCK --
        frame_search = tk.LabelFrame(frame_middle, text="1. Buscar en Inventario", font=("Segoe UI", 10, "bold"),
                                     bg=COLOR_FONDO_MAIN, fg="#2980b9")
        frame_search.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))

        # Buscador Input
        frame_input_search = tk.Frame(frame_search, bg=COLOR_FONDO_MAIN)
        frame_input_search.pack(fill=tk.X, padx=5, pady=5)
        entry_buscar_modal = ttk.Entry(frame_input_search)
        entry_buscar_modal.pack(side=tk.LEFT, fill=tk.X, expand=True)

        # Tabla de Resultados de Búsqueda
        cols_s = ("Codigo", "Nombre", "Stock")
        tree_search = ttk.Treeview(frame_search, columns=cols_s, show='headings')
        tree_search.heading("Codigo", text="Cód")
        tree_search.column("Codigo", width=40, anchor=tk.CENTER)
        tree_search.heading("Nombre", text="Pieza")
        tree_search.column("Nombre", width=120)
        tree_search.heading("Stock", text="Stock")
        tree_search.column("Stock", width=40, anchor=tk.CENTER)

        scroll_s = tk.Scrollbar(frame_search, command=tree_search.yview)
        tree_search.configure(yscrollcommand=scroll_s.set)
        tree_search.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5, pady=5)
        scroll_s.pack(side=tk.RIGHT, fill=tk.Y, pady=5)

        # Lógica de Búsqueda
        def buscar_en_modal(event=None):
            term = entry_buscar_modal.get().strip().lower()
            for i in tree_search.get_children(): tree_search.delete(i)
            for pid, d in self.inventario.items():
                cod = str(d.get('codigo', '')).lower()
                nom = str(d['nombre']).lower()
                if term in cod or term in nom:
                    tree_search.insert("", tk.END, values=(d.get('codigo', ''), d['nombre'], d['cantidad']))

        btn_b_modal = tk.Button(frame_input_search, text="🔍", command=buscar_en_modal,
                                bg="#34495e", fg="white", width=3)
        btn_b_modal.pack(side=tk.LEFT, padx=2)
        entry_buscar_modal.bind('<Return>', buscar_en_modal)

        # -- SUB-PANEL DERECHO: AGREGAR AL CARRITO --
        frame_add = tk.LabelFrame(frame_middle, text="2. Seleccionar", font=("Segoe UI", 10, "bold"),
                                  bg=COLOR_FONDO_MAIN, fg="#e67e22")
        frame_add.pack(side=tk.RIGHT, fill=tk.BOTH, padx=(5, 0))

        tk.Label(frame_add, text="Cód. Seleccionado:", bg=COLOR_FONDO_MAIN).pack(pady=(10, 2))
        entry_cod = ttk.Entry(frame_add, font=("Segoe UI", 11), justify="center", width=15)
        entry_cod.pack(pady=2, padx=10)

        tk.Label(frame_add, text="Cantidad:", bg=COLOR_FONDO_MAIN).pack(pady=(10, 2))
        entry_cant = ttk.Entry(frame_add, font=("Segoe UI", 11), justify="center", width=8)
        entry_cant.insert(0, "1")
        entry_cant.pack(pady=2, padx=10)

        # Función de Búsqueda ID
        def buscar_id_por_codigo(codigo):
            codigo = codigo.strip().lower()
            for pid, data in self.inventario.items():
                if str(data.get('codigo', '')).strip().lower() == codigo:
                    return pid
            return None

        # Función Agregar al Carrito
        def agregar_al_carrito():
            cod = entry_cod.get().strip()
            cant_str = entry_cant.get().strip()

            if not cod: return messagebox.showwarning("Error", "Selecciona una pieza.")
            if not cant_str.isdigit(): return messagebox.showwarning("Error", "Cantidad inválida.")
            cant = int(cant_str)
            if cant <= 0: return messagebox.showwarning("Error", "Cantidad > 0.")

            pid = buscar_id_por_codigo(cod)

            if not pid:
                return messagebox.showerror("Error", f"Código '{cod}' no existe.")

            nombre = self.inventario[pid]['nombre']

            carrito.append({'id': pid, 'codigo': cod, 'nombre': nombre, 'cant': cant})
            actualizar_tabla_carrito()

            entry_cod.delete(0, tk.END)
            entry_cant.delete(0, tk.END)
            entry_cant.insert(0, "1")
            entry_cod.focus()

        # Acción al seleccionar
        def al_seleccionar_busqueda(event):
            sel = tree_search.selection()
            if not sel: return
            vals = tree_search.item(sel[0], 'values')
            entry_cod.delete(0, tk.END)
            entry_cod.insert(0, vals[0])

        tree_search.bind("<<TreeviewSelect>>", al_seleccionar_busqueda)

        def al_doble_clic(event):
            al_seleccionar_busqueda(event)
            entry_cant.focus()

        tree_search.bind("<Double-1>", al_doble_clic)

        btn_add = tk.Button(frame_add, text="⬇️ AGREGAR", command=agregar_al_carrito,
                            bg="#2980b9", fg="white", font=("Segoe UI", 9, "bold"), height=2)
        btn_add.pack(fill=tk.X, padx=10, pady=15)

        entry_cant.bind('<Return>', lambda e: agregar_al_carrito())

        # --- SECCIÓN 3: TABLA CARRITO (Abajo) ---
        tk.Label(ventana, text="3. Carrito de Salida / Entrada", font=("Segoe UI", 10, "bold"),
                 bg=COLOR_FONDO_MAIN, fg="#27ae60").pack(anchor="w", padx=20, pady=(5, 0))

        frame_tabla = tk.Frame(ventana, bg=COLOR_FONDO_MAIN)
        frame_tabla.pack(fill=tk.BOTH, expand=True, padx=20, pady=5)

        cols = ("Codigo", "Nombre", "Cantidad")
        tree_cart = ttk.Treeview(frame_tabla, columns=cols, show='headings', height=5)
        tree_cart.heading("Codigo", text="Código")
        tree_cart.column("Codigo", width=80, anchor=tk.CENTER)
        tree_cart.heading("Nombre", text="Herramienta")
        tree_cart.column("Nombre", width=300)
        tree_cart.heading("Cantidad", text="Cant.")
        tree_cart.column("Cantidad", width=80, anchor=tk.CENTER)

        tree_cart.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scroll = tk.Scrollbar(frame_tabla, command=tree_cart.yview)
        scroll.pack(side=tk.RIGHT, fill=tk.Y)
        tree_cart.config(yscrollcommand=scroll.set)

        def actualizar_tabla_carrito():
            for i in tree_cart.get_children(): tree_cart.delete(i)
            for item in carrito:
                tree_cart.insert("", tk.END, values=(item['codigo'], item['nombre'], item['cant']))

        def limpiar_carrito():
            carrito.clear()
            actualizar_tabla_carrito()

        tk.Button(ventana, text="Limpiar Lista", command=limpiar_carrito,
                  bg="#95a5a6", fg="white", font=("Segoe UI", 8)).pack(anchor="e", padx=20)

        # --- BOTONES DE ACCIÓN FINAL ---
        frame_final = tk.Frame(ventana, bg=COLOR_FONDO_MAIN)
        frame_final.pack(fill=tk.X, padx=20, pady=10)

        def procesar_transaccion(tipo):
            emp = entry_empleado.get().strip()
            if not emp: return messagebox.showerror("Falta Empleado", "Ingresa el Número de Empleado.")
            if not carrito: return messagebox.showwarning("Vacío", "La lista está vacía.")

            accion_txt = "PRESTAR" if tipo == "SALIDA" else "DEVOLVER"

            # Validar stock solo para salidas
            if tipo == "SALIDA":
                for item in carrito:
                    stock_actual = int(self.inventario[item['id']]['cantidad'])
                    if item['cant'] > stock_actual:
                        return messagebox.showerror("Error de Stock",
                                                    f"No hay suficiente stock de:\n{item['nombre']}\n\nStock actual: {stock_actual}\nSolicitado: {item['cant']}")

            # --- MENSAJE DE CONFIRMACIÓN DETALLADO ---
            nombres_lista = ""
            for item in carrito:
                nombres_lista += f"• {item['cant']} pz - {item['nombre']}\n"

            msg_confirm = f"¿Está seguro de {accion_txt} estas herramientas al empleado {emp}?\n\n{nombres_lista}"

            if not messagebox.askyesno("Confirmar Transacción", msg_confirm):
                return
            # -----------------------------------------

            for item in carrito:
                pid = item['id']
                cant = item['cant']
                nombre_pieza = item['nombre']
                codigo_pieza = item['codigo']

                stock_actual = int(self.inventario[pid]['cantidad'])

                if tipo == "SALIDA":
                    nuevo_stock = stock_actual - cant
                    txt_log = f"Préstamo a Empleado: {emp} (-{cant})"
                    msg_historial = f"📤 {datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')} | SALIDA | Empleado: {emp} | Cant: -{cant} | Restan: {nuevo_stock}"
                else:
                    nuevo_stock = stock_actual + cant
                    txt_log = f"Devolución de Empleado: {emp} (+{cant})"
                    msg_historial = f"📥 {datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')} | ENTRADA | Empleado: {emp} | Cant: +{cant} | Total: {nuevo_stock}"

                registrar_accion_global(tipo, codigo_pieza, nombre_pieza, txt_log)
                self.actualizar_inventario_movimiento(pid, nuevo_stock, msg_historial)

            messagebox.showinfo("Éxito", f"Operación completada con {len(carrito)} ítems.")
            ventana.destroy()

        btn_prestar = tk.Button(frame_final, text="📤 PROCESAR PRÉSTAMO",
                                command=lambda: procesar_transaccion("SALIDA"),
                                bg="#e74c3c", fg="white", font=("Segoe UI", 10, "bold"), height=2)
        btn_prestar.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)

        btn_devolver = tk.Button(frame_final, text="📥 PROCESAR DEVOLUCIÓN",
                                 command=lambda: procesar_transaccion("ENTRADA"),
                                 bg="#27ae60", fg="white", font=("Segoe UI", 10, "bold"), height=2)
        btn_devolver.pack(side=tk.RIGHT, fill=tk.X, expand=True, padx=5)

    def actualizar_inventario_movimiento(self, pid, nuevo_stock, mensaje):
        self.inventario[pid]['cantidad'] = nuevo_stock
        if 'historial' not in self.inventario[pid]: self.inventario[pid]['historial'] = []
        self.inventario[pid]['historial'].insert(0, mensaje)
        guardar_datos(self.inventario)
        self.refrescar_tabla(id_seleccionado=pid)

    def obtener_estatus_hoy_texto(self, historial):
        if not historial: return "✅ Sin cambios hoy"
        try:
            ultimo_mov = historial[0]
            partes = ultimo_mov.split(" ")
            if len(partes) > 2:
                fecha_mov = partes[1]
                hora_mov = partes[2][:5]
                hoy = datetime.datetime.now().strftime("%d/%m/%Y")
                if fecha_mov == hoy:
                    if "SALIDA" in ultimo_mov:
                        return f"📤 SALIDA a las {hora_mov}"
                    else:
                        return f"📥 ENTRADA a las {hora_mov}"
        except:
            return "✅ Sin cambios hoy"
        return "✅ Sin cambios hoy"

    def salir_sistema(self):
        if self.verificar_bloqueo(): return
        if messagebox.askyesno("Salir", "¿Deseas cerrar?"): self.root.destroy()

    def cambiar_orden(self, orden):
        if self.verificar_bloqueo(): return
        self.orden_actual = orden;
        self.actualizar_estilo_botones_sort();
        self.refrescar_tabla(self.obtener_id_actual_seleccionado())

    # --- AQUÍ ESTÁ LA FUNCIÓN ABRIR_AYUDA MEJORADA ---
    def abrir_ayuda(self):
        if self.verificar_bloqueo(): return

        ventana = tk.Toplevel(self.root)
        ventana.title("Manual de Usuario y Documentación v2.0")
        ventana.geometry("900x750")
        ventana.configure(bg="#ecf0f1")

        # --- Encabezado Fijo ---
        frame_head = tk.Frame(ventana, bg="#2c3e50", height=60)
        frame_head.pack(fill=tk.X)
        frame_head.pack_propagate(False)
        tk.Label(frame_head, text="📘 MANUAL DE OPERACIONES | SISTEMA DE INVENTARIO",
                 font=("Segoe UI", 16, "bold"), bg="#2c3e50", fg="white").pack(pady=15)

        # --- Área de Texto con Scroll ---
        frame_txt = tk.Frame(ventana, bg="white")
        frame_txt.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        scroll = tk.Scrollbar(frame_txt)
        scroll.pack(side=tk.RIGHT, fill=tk.Y)

        txt = tk.Text(frame_txt, font=("Segoe UI", 10), bg="white", fg="#2c3e50",
                      relief="flat", wrap="word", yscrollcommand=scroll.set, padx=20, pady=20)
        txt.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scroll.config(command=txt.yview)

        # --- CONFIGURACIÓN DE ESTILOS (TAGS) ---
        # Títulos de secciones
        txt.tag_config("h1", font=("Segoe UI", 14, "bold"), foreground="#2980b9", spacing1=20, spacing3=10)
        # Subtítulos
        txt.tag_config("h2", font=("Segoe UI", 11, "bold"), foreground="#34495e", spacing1=10)
        # Texto normal
        txt.tag_config("normal", font=("Segoe UI", 10), spacing1=2)
        # Resaltado de botones (simula un botón visualmente)
        txt.tag_config("btn", font=("Segoe UI", 9, "bold"), background="#ecf0f1", foreground="#c0392b")
        # Alertas o notas importantes
        txt.tag_config("nota", font=("Segoe UI", 9, "italic"), foreground="#7f8c8d", background="#fff3cd", lmargin1=20)
        # Listas
        txt.tag_config("bullet", lmargin1=20, lmargin2=30, spacing1=3)

        # --- CONTENIDO DEL MANUAL ---

        # 1. INTRODUCCIÓN
        txt.insert(tk.END, "1. VISIÓN GENERAL DEL SISTEMA\n", "h1")
        txt.insert(tk.END,
                   "Bienvenido al Sistema de Gestión de Inventario. Este software permite controlar el flujo de herramientas, gestionar préstamos a empleados y generar auditorías automáticas.\n",
                   "normal")

        # 2. GESTIÓN DE INVENTARIO
        txt.insert(tk.END, "2. GESTIÓN DE INVENTARIO (Panel Izquierdo)\n", "h1")
        txt.insert(tk.END, "Aquí se administra la base de datos de herramientas.\n\n", "normal")

        txt.insert(tk.END, "• Alta de Material:\n", "h2")
        txt.insert(tk.END, "   1. Presione ", "normal");
        txt.insert(tk.END, "[✨ NUEVO / LIMPIAR]", "btn");
        txt.insert(tk.END, ".\n", "normal")
        txt.insert(tk.END, "   2. Llene los campos (Código, Nombre, Stock, etc.).\n", "normal")
        txt.insert(tk.END, "   3. Presione ", "normal");
        txt.insert(tk.END, "[💾 GUARDAR NUEVO]", "btn");
        txt.insert(tk.END, " para confirmar.\n", "normal")

        txt.insert(tk.END, "\n• Modificación:\n", "h2")
        txt.insert(tk.END, "Seleccione una pieza de la tabla, presione ", "normal");
        txt.insert(tk.END, "[✏️ MODIFICAR]", "btn");
        txt.insert(tk.END, ", edite los datos y guarde.\n", "normal")

        txt.insert(tk.END, "\n• Semáforo de Stock (NUEVO):\n", "h2")
        txt.insert(tk.END, "El sistema le avisará con colores cuando se acaben las piezas:\n", "normal")
        txt.insert(tk.END, "   🔴 ROJO: Stock Crítico (Quedan 2 o menos).\n", "bullet")
        txt.insert(tk.END, "   🟡 AMARILLO: Stock Bajo (Quedan 5 o menos).\n", "bullet")

        txt.insert(tk.END,
                   "\n⚠️ NOTA: El código de la herramienta debe ser único. El sistema no permitirá duplicados.\n",
                   "nota")

        # 3. PRÉSTAMOS
        txt.insert(tk.END, "3. MÓDULO DE PRÉSTAMOS (Botón Morado)\n", "h1")
        txt.insert(tk.END, "Esta es la función principal para el día a día.\n\n", "normal")

        txt.insert(tk.END, "PASO 1: Identificación\n", "h2")
        txt.insert(tk.END, "Ingrese el NÚMERO DE EMPLEADO en la parte superior. \n", "normal")
        txt.insert(tk.END, "⚠️ El sistema bloqueará letras automáticamente, solo ingrese números.\n", "nota")

        txt.insert(tk.END, "\nPASO 2: Armado del Carrito\n", "h2")
        txt.insert(tk.END, "Busque la herramienta en la lista izquierda, selecciónela e indique la cantidad. Presione ",
                   "normal");
        txt.insert(tk.END, "[⬇️ AGREGAR]", "btn");
        txt.insert(tk.END, ".\n", "normal")
        txt.insert(tk.END, "Repita esto para todas las herramientas que el empleado necesite.\n", "bullet")

        txt.insert(tk.END, "\nPASO 3: Confirmación\n", "h2")
        txt.insert(tk.END, "• Para entregar material (Resta Stock): Use ", "normal");
        txt.insert(tk.END, "[📤 PROCESAR PRÉSTAMO]", "btn");
        txt.insert(tk.END, ".\n", "normal")
        txt.insert(tk.END, "• Para recibir material (Suma Stock): Use ", "normal");
        txt.insert(tk.END, "[📥 PROCESAR DEVOLUCIÓN]", "btn");
        txt.insert(tk.END, ".\n", "normal")

        # 4. REPORTES
        txt.insert(tk.END, "4. REPORTES Y AUDITORÍA\n", "h1")
        txt.insert(tk.END, "Acceda desde el botón azul [📂 CENTRO DE REPORTES].\n", "normal")

        txt.insert(tk.END, "• Reporte del Día:", "h2");
        txt.insert(tk.END, " Muestra cronológicamente qué pasó hoy (quién sacó qué). Ideal para el cierre de turno.\n",
                   "bullet")
        txt.insert(tk.END, "• Estatus Empleados:", "h2");
        txt.insert(tk.END,
                   " Muestra una lista de quién debe material. Si está en ROJO, el empleado no ha devuelto la herramienta.\n",
                   "bullet")
        txt.insert(tk.END, "• Exportación:", "h2");
        txt.insert(tk.END, " Puede descargar todo el inventario en Excel, Word o PDF para inventarios físicos.\n",
                   "bullet")

        # 5. SOLUCIÓN DE PROBLEMAS
        txt.insert(tk.END, "5. PREGUNTAS FRECUENTES (FAQ)\n", "h1")

        txt.insert(tk.END, "¿Por qué no me deja borrar una pieza?\n", "h2")
        txt.insert(tk.END,
                   "Asegúrese de haber seleccionado la fila en la tabla principal. El campo 'ID' debe tener un número.\n",
                   "bullet")

        txt.insert(tk.END, "\n¿Dónde se guardan los respaldos?\n", "h2")
        txt.insert(tk.END,
                   "En la carpeta 'Respaldos' ubicada junto al ejecutable del sistema. Se recomienda copiar esta carpeta a una USB periódicamente.\n",
                   "bullet")

        txt.insert(tk.END, "\nEl reporte de Excel da error al guardar.\n", "h2")
        txt.insert(tk.END,
                   "Verifique que no tenga el archivo de Excel abierto. Debe cerrarlo antes de generar uno nuevo con el mismo nombre.\n",
                   "bullet")

        txt.config(state=tk.DISABLED)

        # Botón de cierre
        btn_close = tk.Button(ventana, text="Entendido, cerrar manual", command=ventana.destroy,
                              bg="#34495e", fg="white", font=("Segoe UI", 10), pady=8)
        btn_close.pack(fill=tk.X, padx=20, pady=(0, 20))

    # --- NUEVA FUNCIONALIDAD: REPORTE EMPLEADOS (AGRUPADO) ---
    def ver_estatus_prestamos_empleados(self):
        log = cargar_log_global()
        hoy = datetime.datetime.now().strftime("%d/%m/%Y")

        movimientos_hoy = [m for m in log if m['fecha'] == hoy]

        if not movimientos_hoy:
            return messagebox.showinfo("Reporte Empleados", "No hay movimientos registrados hoy.")

        # Diccionario para agrupar por (Empleado, Codigo)
        # Clave: (emp, cod) -> Valor: {nombre, sacados, devueltos}
        reporte_agregado = {}

        for m in movimientos_hoy:

            # FILTRO ESTRICTO (SOLO PRÉSTAMOS Y DEVOLUCIONES)
            if m['accion'] not in ["SALIDA", "ENTRADA"]:
                continue

            detalle = m['detalle']
            codigo = m['codigo']
            nombre = m['nombre']
            accion = m['accion']

            empleado = "Desconocido"

            # LOGICA DE LIMPIEZA DE NOMBRE (Version v85.0)
            try:
                partes = detalle.split(" ")
                # Buscar directamente el tag "Empleado:" y tomar el siguiente valor
                if "Empleado:" in partes:
                    idx = partes.index("Empleado:")
                    if len(partes) > idx + 1:
                        raw_emp = partes[idx + 1]
                        # Limpiar si por error se guardó "Empleado:123" pegado
                        empleado = raw_emp.replace("Empleado:", "").strip()

                # Compatibilidad con formatos anteriores
                elif "Préstamo" in detalle and "a" in partes:
                    idx = partes.index("a")
                    if len(partes) > idx + 1: empleado = partes[idx + 1]
                elif "Devolución" in detalle and "de" in partes:
                    idx = partes.index("de")
                    if len(partes) > idx + 1: empleado = partes[idx + 1]
            except:
                continue

            cantidad = 0
            try:
                if "(" in detalle and ")" in detalle:
                    inicio = detalle.rfind("(") + 1
                    fin = detalle.rfind(")")
                    cant_str = detalle[inicio:fin]
                    cantidad = int(cant_str.replace("+", "").replace("-", ""))
                else:
                    cantidad = 1
            except:
                cantidad = 1

            key = (empleado, codigo)
            if key not in reporte_agregado:
                reporte_agregado[key] = {'nombre': nombre, 'sacados': 0, 'devueltos': 0}

            if accion == "SALIDA":
                reporte_agregado[key]['sacados'] += cantidad
            elif accion == "ENTRADA":
                reporte_agregado[key]['devueltos'] += cantidad

        ventana = tk.Toplevel(self.root)
        ventana.title(f"Estatus de Empleados - {hoy}")
        ventana.geometry("1000x500")
        ventana.configure(bg=COLOR_FONDO_MAIN)
        ventana.grab_set()

        lbl = tk.Label(ventana, text="CONTROL DE PRÉSTAMOS POR EMPLEADO (HOY)",
                       font=("Segoe UI", 12, "bold"), bg=COLOR_FONDO_MAIN, fg=COLOR_TEXTO_OSCURO)
        lbl.pack(pady=10)

        cols = ("Empleado", "Codigo", "Pieza", "Sacados", "Devueltos", "Pendiente", "Estatus")
        tree = ttk.Treeview(ventana, columns=cols, show='headings')

        tree.heading("Empleado", text="Empleado")
        tree.column("Empleado", width=120, anchor=tk.CENTER)

        tree.heading("Codigo", text="Cód.")
        tree.column("Codigo", width=70, anchor=tk.CENTER)
        tree.heading("Pieza", text="Herramienta")
        tree.column("Pieza", width=250)
        tree.heading("Sacados", text="Sacados")
        tree.column("Sacados", width=70, anchor=tk.CENTER)
        tree.heading("Devueltos", text="Devueltos")
        tree.column("Devueltos", width=70, anchor=tk.CENTER)
        tree.heading("Pendiente", text="Debe")
        tree.column("Pendiente", width=70, anchor=tk.CENTER)
        tree.heading("Estatus", text="Estatus")
        tree.column("Estatus", width=150, anchor=tk.CENTER)

        tree.tag_configure('pendiente', foreground='red')
        tree.tag_configure('parcial', foreground='#d35400')  # Naranja
        tree.tag_configure('devuelto', foreground='green')

        pendientes_count = 0

        for (emp, cod), data in reporte_agregado.items():
            sacados = data['sacados']
            devueltos = data['devueltos']
            pendiente = sacados - devueltos

            if pendiente == 0:
                estatus_txt = "DEVUELTO TOTAL ✅"
                tag = 'devuelto'
            elif pendiente > 0 and devueltos > 0:
                estatus_txt = f"PARCIAL (Faltan {pendiente})"
                tag = 'parcial'
                pendientes_count += 1
            elif pendiente == sacados:
                estatus_txt = "PENDIENTE ❌"
                tag = 'pendiente'
                pendientes_count += 1
            else:
                estatus_txt = "SALDO A FAVOR"
                tag = 'devuelto'

            # SOLO PONEMOS EL NUMERO DEL EMPLEADO (SIN TEXTO EXTRA)
            emp_str = str(emp).replace("Empleado:", "").strip()

            tree.insert("", tk.END, values=(
                emp_str, cod, data['nombre'],
                sacados, devueltos, pendiente, estatus_txt
            ), tags=(tag,))

        tree.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

        resumen_txt = f"Registros Activos: {len(reporte_agregado)} | Empleados con Pendientes: {pendientes_count}"
        lbl_res = tk.Label(ventana, text=resumen_txt, font=("Segoe UI", 10, "bold"),
                           bg="#ecf0f1", fg="#c0392b" if pendientes_count > 0 else "#27ae60")
        lbl_res.pack(pady=10)

        tk.Button(ventana, text="Cerrar", command=ventana.destroy, bg=COLOR_BTN_SALIR, fg="white").pack(pady=5)

    # --- CENTRO DE REPORTES ---
    def abrir_centro_reportes(self):
        if self.verificar_bloqueo(): return
        if not self.inventario: return messagebox.showwarning("Vacío", "No hay inventario.")
        pid = self.obtener_id_actual_seleccionado()

        ventana = tk.Toplevel(self.root)
        ventana.title("Centro de Reportes")
        ventana.geometry("500x750")
        ventana.configure(bg=COLOR_FONDO_MAIN)
        ventana.grab_set()

        font_titulo = ("Segoe UI", 11, "bold")

        tk.Label(ventana, text="📊 Reportes de Inventario (Existencias)", font=font_titulo, bg=COLOR_FONDO_MAIN,
                 fg=COLOR_TEXTO_OSCURO).pack(pady=(20, 5))
        frame_inv = tk.Frame(ventana, bg=COLOR_FONDO_MAIN, bd=1, relief="solid")
        frame_inv.pack(fill=tk.X, padx=20, pady=5)

        def exp_excel_todo():
            if messagebox.askyesno("Confirmar", "¿Exportar TODO el inventario a Excel?"):
                ventana.destroy();
                self.generar_excel_general()

        def exp_word_todo():
            if messagebox.askyesno("Confirmar", "¿Exportar TODO el inventario a Word?"):
                ventana.destroy();
                self.generar_word_general()

        def exp_pdf_todo():
            if messagebox.askyesno("Confirmar", "¿Exportar TODO el inventario a PDF?"):
                ventana.destroy();
                self.generar_pdf_general()

        def exp_excel_uno():
            if not pid: return messagebox.showwarning("Ojo", "Selecciona una pieza primero.")
            if messagebox.askyesno("Confirmar", "Exportar ficha a Excel?"):
                ventana.destroy();
                self.generar_excel_individual(pid)

        def exp_word_uno():
            if not pid: return messagebox.showwarning("Ojo", "Selecciona una pieza primero.")
            if messagebox.askyesno("Confirmar", "Exportar ficha a Word?"):
                ventana.destroy();
                self.generar_word_individual(pid)

        def exp_pdf_uno():
            if not pid: return messagebox.showwarning("Ojo", "Selecciona una pieza primero.")
            if messagebox.askyesno("Confirmar", "Exportar ficha a PDF?"):
                ventana.destroy();
                self.generar_pdf_individual(pid)

        tk.Button(frame_inv, text="Excel (Todo)", command=exp_excel_todo, bg=COLOR_BTN_EXCEL, fg="white", relief="flat",
                  cursor="hand2", image=self.pixel, compound="center", height=25, width=130).pack(side=tk.LEFT, padx=5,
                                                                                                  pady=10)
        tk.Button(frame_inv, text="Word (Todo)", command=exp_word_todo, bg=COLOR_BTN_WORD, fg="white", relief="flat",
                  cursor="hand2", image=self.pixel, compound="center", height=25, width=130).pack(side=tk.LEFT, padx=5,
                                                                                                  pady=10)
        tk.Button(frame_inv, text="PDF (Todo)", command=exp_pdf_todo, bg=COLOR_BTN_PDF, fg="white", relief="flat",
                  cursor="hand2", image=self.pixel, compound="center", height=25, width=130).pack(side=tk.LEFT, padx=5,
                                                                                                  pady=10)

        if pid:
            nombre_sel = self.inventario[pid]['nombre'][:25] + "..."
            frame_ind = tk.Frame(ventana, bg=COLOR_HIGHLIGHT_BOX, bd=2, relief="groove")
            frame_ind.pack(fill=tk.X, padx=20, pady=10)

            tk.Label(frame_ind, text=f"✅ HERRAMIENTA SELECCIONADA: {nombre_sel}",
                     bg=COLOR_HIGHLIGHT_BOX, fg="black", font=("Segoe UI", 9, "bold")).pack(pady=(10, 5))

            frame_btns_ind = tk.Frame(frame_ind, bg=COLOR_HIGHLIGHT_BOX)
            frame_btns_ind.pack(pady=10)

            tk.Button(frame_btns_ind, text="Excel (Pieza)", command=exp_excel_uno, bg=COLOR_BTN_EXCEL, fg="white",
                      relief="flat", cursor="hand2", image=self.pixel, compound="center", height=25, width=100).pack(
                side=tk.LEFT, padx=5)
            tk.Button(frame_btns_ind, text="Word (Pieza)", command=exp_word_uno, bg=COLOR_BTN_WORD, fg="white",
                      relief="flat", cursor="hand2", image=self.pixel, compound="center", height=25, width=100).pack(
                side=tk.LEFT, padx=5)
            tk.Button(frame_btns_ind, text="PDF (Pieza)", command=exp_pdf_uno, bg=COLOR_BTN_PDF, fg="white",
                      relief="flat", cursor="hand2", image=self.pixel, compound="center", height=25, width=100).pack(
                side=tk.LEFT, padx=5)

        tk.Label(ventana, text="📅 Auditoría del Día (Movimientos)", font=font_titulo, bg=COLOR_FONDO_MAIN,
                 fg=COLOR_TEXTO_OSCURO).pack(pady=(20, 5))
        frame_dia = tk.Frame(ventana, bg=COLOR_FONDO_MAIN, bd=1, relief="solid")
        frame_dia.pack(fill=tk.X, padx=20, pady=5)

        def reporte_dia_excel():
            if not messagebox.askyesno("Confirmar", "¿Desea generar el reporte de auditoría del día en EXCEL?"): return
            ventana.destroy();
            self.generar_reporte_dia("Excel")

        def reporte_dia_word():
            if not messagebox.askyesno("Confirmar", "¿Desea generar el reporte de auditoría del día en WORD?"): return
            ventana.destroy();
            self.generar_reporte_dia("Word")

        def reporte_dia_pdf():
            if not messagebox.askyesno("Confirmar", "¿Desea generar el reporte de auditoría del día en PDF?"): return
            ventana.destroy();
            self.generar_reporte_dia_pdf()

        def reporte_pantalla():
            self.ver_reporte_pantalla()

        def reporte_empleados():
            self.ver_estatus_prestamos_empleados()

        tk.Button(frame_dia, text="👁️ VER EN PANTALLA (Historial)", command=reporte_pantalla, bg=COLOR_BTN_VISOR,
                  fg="white",
                  relief="flat", cursor="hand2", image=self.pixel, compound="center", height=25, width=250,
                  font=("Segoe UI", 9, "bold")).pack(fill=tk.X, padx=10, pady=2)

        tk.Button(frame_dia, text="👷 ESTATUS PRÉSTAMOS EMPLEADOS", command=reporte_empleados,
                  bg=COLOR_BTN_EMPLEADOS, fg="white", relief="flat", cursor="hand2",
                  image=self.pixel, compound="center", height=30, width=250,
                  font=("Segoe UI", 9, "bold")).pack(fill=tk.X, padx=10, pady=5)

        frame_dia_btns = tk.Frame(frame_dia, bg=COLOR_FONDO_MAIN)
        frame_dia_btns.pack(pady=5)
        tk.Button(frame_dia_btns, text="Excel", command=reporte_dia_excel, bg=COLOR_BTN_EXCEL, fg="white", height=25,
                  width=100, image=self.pixel, compound="center").pack(side=tk.LEFT, padx=2)
        tk.Button(frame_dia_btns, text="Word", command=reporte_dia_word, bg=COLOR_BTN_WORD, fg="white", height=25,
                  width=100, image=self.pixel, compound="center").pack(side=tk.LEFT, padx=2)
        tk.Button(frame_dia_btns, text="PDF", command=reporte_dia_pdf, bg=COLOR_BTN_PDF, fg="white", height=25,
                  width=100, image=self.pixel, compound="center").pack(side=tk.LEFT, padx=2)

        tk.Label(ventana, text="🛡️ Seguridad de Datos", font=font_titulo, bg=COLOR_FONDO_MAIN,
                 fg=COLOR_TEXTO_OSCURO).pack(pady=(20, 5))
        frame_sec = tk.Frame(ventana, bg=COLOR_FONDO_MAIN, bd=1, relief="solid")
        frame_sec.pack(fill=tk.X, padx=20, pady=5)
        tk.Button(frame_sec, text="💾 CREAR RESPALDO AHORA", command=self.crear_respaldo_manual, bg=COLOR_BTN_RESPALDO,
                  fg="white", relief="flat", cursor="hand2", image=self.pixel, compound="center", height=30, width=200,
                  font=("Segoe UI", 8, "bold")).pack(side=tk.LEFT, padx=5, pady=5)

        tk.Button(frame_sec, text="📂 LEER RESPALDO ANTERIOR", command=self.abrir_visor_respaldos,
                  bg=COLOR_BTN_LEER_RESPALDO, fg="white", relief="flat", cursor="hand2", image=self.pixel,
                  compound="center", height=30, width=200, font=("Segoe UI", 8, "bold")).pack(side=tk.LEFT, padx=5,
                                                                                              pady=5)

        tk.Button(ventana, text="Cancelar", command=ventana.destroy, bg="#7f8c8d", fg="white", relief="flat",
                  cursor="hand2", image=self.pixel, compound="center", height=25, width=100).pack(pady=15)

    def ver_reporte_pantalla(self):
        log = cargar_log_global()
        hoy = datetime.datetime.now().strftime("%d/%m/%Y")
        movimientos_hoy = [m for m in log if m['fecha'] == hoy]

        if not movimientos_hoy: return messagebox.showinfo("Aviso", "No hay movimientos hoy.")

        movimientos_hoy.sort(key=lambda x: x['hora'], reverse=True)

        ventana = tk.Toplevel(self.root)
        ventana.title(f"Actividad en Tiempo Real - {hoy}")
        ventana.geometry("800x600")
        ventana.configure(bg=COLOR_FONDO_MAIN)
        ventana.grab_set()

        resumen = {"CREACIÓN": 0, "ELIMINACIÓN": 0, "SALIDA": 0, "ENTRADA": 0}
        for m in movimientos_hoy:
            if m['accion'] in resumen: resumen[m['accion']] += 1

        frame_resumen = tk.Frame(ventana, bg="white", padx=20, pady=10)
        frame_resumen.pack(fill=tk.X, padx=20, pady=10)

        tk.Label(frame_resumen, text=f"Salidas: {resumen['SALIDA']}", fg="blue", font=("Segoe UI", 10, "bold"),
                 bg="white").pack(side=tk.LEFT, padx=10)
        tk.Label(frame_resumen, text=f"Entradas: {resumen['ENTRADA']}", fg="green", font=("Segoe UI", 10, "bold"),
                 bg="white").pack(side=tk.LEFT, padx=10)
        tk.Label(frame_resumen, text=f"Nuevas: {resumen['CREACIÓN']}", fg="black", font=("Segoe UI", 10, "bold"),
                 bg="white").pack(side=tk.LEFT, padx=10)
        tk.Label(frame_resumen, text=f"Bajas: {resumen['ELIMINACIÓN']}", fg="red", font=("Segoe UI", 10, "bold"),
                 bg="white").pack(side=tk.LEFT, padx=10)

        cols = ("Hora", "Accion", "Codigo", "Pieza", "Detalle")
        tree = ttk.Treeview(ventana, columns=cols, show='headings')
        tree.heading("Hora", text="Hora");
        tree.column("Hora", width=80, anchor=tk.CENTER)
        tree.heading("Accion", text="Acción");
        tree.column("Accion", width=100, anchor=tk.CENTER)
        tree.heading("Codigo", text="Código");
        tree.column("Codigo", width=80, anchor=tk.CENTER)
        tree.heading("Pieza", text="Pieza");
        tree.column("Pieza", width=200)
        tree.heading("Detalle", text="Detalle");
        tree.column("Detalle", width=300)

        tree.tag_configure("SALIDA", foreground="blue")
        tree.tag_configure("ENTRADA", foreground="green")
        tree.tag_configure("ELIMINACIÓN", foreground="red")

        for m in movimientos_hoy:
            tree.insert("", tk.END, values=(m['hora'], m['accion'], m['codigo'], m['nombre'], m['detalle']),
                        tags=(m['accion'],))

        tree.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 10))

        btn_continuar = tk.Button(ventana, text="Continuar", command=ventana.destroy,
                                  bg="#34495e", fg="white", font=("Segoe UI", 9, "bold"),
                                  relief="flat", cursor="hand2", image=self.pixel, compound="center", height=30,
                                  width=120)
        btn_continuar.pack(pady=10)

    def generar_reporte_dia(self, formato):
        log = cargar_log_global()
        hoy = datetime.datetime.now().strftime("%d/%m/%Y")
        movimientos_hoy = [m for m in log if m['fecha'] == hoy]

        if not movimientos_hoy: return messagebox.showinfo("Reporte Diario", "No hay actividad registrada hoy.")
        movimientos_hoy.sort(key=lambda x: x['hora'])

        resumen = {"CREACIÓN": 0, "ELIMINACIÓN": 0, "SALIDA": 0, "ENTRADA": 0}
        for m in movimientos_hoy:
            if m['accion'] in resumen: resumen[m['accion']] += 1

        nombre_archivo = f"Reporte_Diario_{hoy.replace('/', '-')}"

        if formato == "Excel":
            ruta = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel", "*.xlsx")],
                                                initialfile=f"{nombre_archivo}.xlsx")
            if not ruta: return
            wb = openpyxl.Workbook();
            ws = wb.active;
            ws.title = "Reporte Diario"
            ws.append([f"REPORTE DE ACTIVIDAD - {hoy}"])
            ws.append([])
            ws.append(["--- RESUMEN ESTADÍSTICO ---"])
            ws.append(["Acción", "Total Eventos"])
            ws.append(["✅ Piezas Creadas", resumen["CREACIÓN"]])
            ws.append(["🗑️ Piezas Eliminadas", resumen["ELIMINACIÓN"]])
            ws.append(["📤 Préstamos (Salidas)", resumen["SALIDA"]])
            ws.append(["📥 Devoluciones (Entradas)", resumen["ENTRADA"]])
            ws.append([])
            ws.append(["HORA", "ACCIÓN", "CÓDIGO", "PIEZA", "DETALLES"])
            for m in movimientos_hoy:
                ws.append([m['hora'], m['accion'], m['codigo'], m['nombre'], m['detalle']])
            wb.save(ruta)
            if messagebox.askyesno("Listo", "¿Abrir reporte?"): os.startfile(ruta)

        else:  # Word
            ruta = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word", "*.docx")],
                                                initialfile=f"{nombre_archivo}.docx")
            if not ruta: return
            doc = Document()
            doc.add_heading(f"REPORTE DIARIO DE ACTIVIDAD", 0)
            doc.add_paragraph(f"Fecha: {hoy}")
            doc.add_heading("1. Resumen Estadístico", level=2)
            tr = doc.add_table(rows=1, cols=2);
            tr.style = 'Table Grid'
            tr.rows[0].cells[0].text = "Acción";
            tr.rows[0].cells[1].text = "Cantidad Total"
            tr.add_row().cells[0].text = "✅ Piezas Creadas";
            tr.rows[1].cells[1].text = str(resumen["CREACIÓN"])
            doc.add_paragraph("")
            doc.add_heading("2. Detalle Cronológico", level=2)
            t = doc.add_table(rows=1, cols=5);
            t.style = 'Table Grid'
            h = t.rows[0].cells;
            h[0].text = "HORA";
            h[1].text = "ACCIÓN";
            h[2].text = "CÓDIGO";
            h[3].text = "PIEZA";
            h[4].text = "DETALLES"
            for m in movimientos_hoy:
                r = t.add_row().cells
                r[0].text = m['hora'];
                r[1].text = m['accion'];
                r[2].text = str(m['codigo']);
                r[3].text = m['nombre'];
                r[4].text = m['detalle']
            doc.save(ruta)
            if messagebox.askyesno("Listo", "¿Abrir reporte?"): os.startfile(ruta)

    def generar_reporte_dia_pdf(self):
        log = cargar_log_global()
        hoy = datetime.datetime.now().strftime("%d/%m/%Y")
        movimientos_hoy = [m for m in log if m['fecha'] == hoy]
        if not movimientos_hoy: return messagebox.showinfo("Aviso", "No hay movimientos hoy.")

        movimientos_hoy.sort(key=lambda x: x['hora'])

        nombre_archivo = f"Reporte_Diario_{hoy.replace('/', '-')}"
        ruta = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF", "*.pdf")],
                                            initialfile=f"{nombre_archivo}.pdf")
        if not ruta: return

        pdf = PDF()
        pdf.add_page()
        pdf.set_font("Arial", size=10)
        pdf.cell(0, 10, f"Fecha del Reporte: {hoy}", 0, 1)
        pdf.ln(5)

        pdf.set_font("Arial", "B", 9)
        pdf.cell(20, 8, "Hora", 1)
        pdf.cell(25, 8, "Accion", 1)
        pdf.cell(20, 8, "Codigo", 1)
        pdf.cell(50, 8, "Pieza", 1)
        pdf.cell(75, 8, "Detalle", 1)
        pdf.ln()

        pdf.set_font("Arial", size=8)
        for m in movimientos_hoy:
            pdf.cell(20, 8, pdf.clean_text(m['hora']), 1)
            pdf.cell(25, 8, pdf.clean_text(m['accion']), 1)
            pdf.cell(20, 8, pdf.clean_text(str(m['codigo'])), 1)
            pdf.cell(50, 8, pdf.clean_text(str(m['nombre'])[:25]), 1)
            pdf.cell(75, 8, pdf.clean_text(str(m['detalle'])[:40]), 1)
            pdf.ln()

        pdf.output(ruta)
        if messagebox.askyesno("Listo", "¿Abrir PDF?"): os.startfile(ruta)

    def generar_excel_general(self):
        ruta = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel", "*.xlsx")],
                                            initialfile="Inventario_Completo.xlsx")
        if not ruta: return
        try:
            wb = openpyxl.Workbook();
            ws = wb.active;
            ws.title = "Inventario"
            ws.append(["ID", "Código", "Nombre", "Cantidad", "Ubicación", "Estatus Hoy", "Descripción"])
            for pid, d in sorted(self.inventario.items(), key=lambda x: int(x[0]) if x[0].isdigit() else 0):
                estatus = self.obtener_estatus_hoy_texto(d.get('historial', []))
                ws.append(
                    [int(pid) if pid.isdigit() else pid, d.get('codigo', ''), d['nombre'], d['cantidad'], d['gabinete'],
                     estatus, d['descripcion']])
            wb.save(ruta)
            if messagebox.askyesno("Éxito", "Reporte generado.\n¿Abrir?"): os.startfile(ruta)
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def generar_word_general(self):
        ruta = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word", "*.docx")],
                                            initialfile="Inventario_Completo.docx")
        if not ruta: return
        try:
            doc = Document();
            doc.add_heading('Inventario', 0)
            t = doc.add_table(rows=1, cols=7);
            t.style = 'Table Grid'
            h = t.rows[0].cells;
            h[0].text = 'ID';
            h[1].text = 'COD';
            h[2].text = 'NOM';
            h[3].text = 'CANT';
            h[4].text = 'UBIC';
            h[5].text = 'EST';
            h[6].text = 'DESC'
            for pid, d in sorted(self.inventario.items(), key=lambda x: int(x[0]) if x[0].isdigit() else 0):
                r = t.add_row().cells;
                est = self.obtener_estatus_hoy_texto(d.get('historial', []))
                r[0].text = str(pid);
                r[1].text = str(d.get('codigo', ''));
                r[2].text = str(d['nombre']);
                r[3].text = str(d['cantidad'])
                r[4].text = str(d['gabinete']);
                r[5].text = est;
                r[6].text = str(d['descripcion'])
            doc.save(ruta);
            if messagebox.askyesno("Éxito", "Reporte generado.\n¿Abrir?"): os.startfile(ruta)
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def generar_pdf_general(self):
        ruta = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF", "*.pdf")],
                                            initialfile="Inventario_Completo.pdf")
        if not ruta: return

        pdf = PDF()
        pdf.add_page()
        pdf.set_font("Arial", "B", 9)
        pdf.cell(10, 8, "ID", 1)
        pdf.cell(20, 8, "Cod", 1)
        pdf.cell(60, 8, "Nombre", 1)
        pdf.cell(15, 8, "Cant", 1)
        pdf.cell(20, 8, "Ubic", 1)
        pdf.cell(65, 8, "Desc", 1)
        pdf.ln()

        pdf.set_font("Arial", size=8)
        for pid, d in sorted(self.inventario.items(), key=lambda x: int(x[0]) if x[0].isdigit() else 0):
            pdf.cell(10, 8, pdf.clean_text(str(pid)), 1)
            pdf.cell(20, 8, pdf.clean_text(str(d.get('codigo', ''))), 1)
            pdf.cell(60, 8, pdf.clean_text(str(d['nombre'])[:30]), 1)
            pdf.cell(15, 8, pdf.clean_text(str(d['cantidad'])), 1)
            pdf.cell(20, 8, pdf.clean_text(str(d['gabinete'])), 1)
            pdf.cell(65, 8, pdf.clean_text(str(d['descripcion'])[:35]), 1)
            pdf.ln()

        pdf.output(ruta)
        if messagebox.askyesno("Éxito", "PDF generado.\n¿Abrir?"): os.startfile(ruta)

    def generar_excel_individual(self, pid):
        d = self.inventario[pid]
        ruta = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel", "*.xlsx")],
                                            initialfile=f"Ficha_{d['nombre']}.xlsx")
        if not ruta: return
        try:
            wb = openpyxl.Workbook();
            ws = wb.active;
            ws.title = "Ficha Técnica"
            ws.append(["FICHA TÉCNICA"]);
            ws.append(["ID", "Código", "Nombre", "Stock", "Ubicación"])
            ws.append(
                [int(pid) if pid.isdigit() else pid, d.get('codigo', ''), d['nombre'], d['cantidad'], d['gabinete']])
            ws.append([]);
            ws.append(["DESCRIPCIÓN:", d['descripcion']]);
            ws.append([]);
            ws.append(["HISTORIAL"])
            for linea in d.get('historial', []): ws.append([linea])
            wb.save(ruta)
            if messagebox.askyesno("Éxito", "Ficha guardada.\n¿Abrir?"): os.startfile(ruta)
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def generar_word_individual(self, pid):
        d = self.inventario[pid]
        ruta = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word", "*.docx")],
                                            initialfile=f"Ficha_{d['nombre']}.docx")
        if not ruta: return
        try:
            doc = Document();
            doc.add_heading(f"Ficha: {d['nombre']}", 0)
            doc.add_paragraph(
                f"ID: {pid} | COD: {d.get('codigo', '')} | STOCK: {d['cantidad']} | UBIC: {d['gabinete']}")
            doc.add_heading("Descripción", 2);
            doc.add_paragraph(d['descripcion'])
            doc.add_heading("Historial", 2)
            if d.get('historial', []):
                t = doc.add_table(rows=1, cols=1);
                t.style = 'Table Grid'
                for l in d['historial']: t.add_row().cells[0].text = l
            else:
                doc.add_paragraph("Sin movimientos.")
            doc.save(ruta)
            if messagebox.askyesno("Éxito", "Ficha guardada.\n¿Abrir?"): os.startfile(ruta)
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def generar_pdf_individual(self, pid):
        d = self.inventario[pid]
        ruta = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF", "*.pdf")],
                                            initialfile=f"Ficha_{d['nombre']}.pdf")
        if not ruta: return

        pdf = PDF()
        pdf.add_page()

        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, pdf.clean_text(f"FICHA TECNICA: {d['nombre']}"), 0, 1)
        pdf.ln(5)

        pdf.set_font("Arial", "B", 10)
        pdf.cell(30, 8, "CODIGO:", 0)
        pdf.set_font("Arial", "", 10)
        pdf.cell(50, 8, pdf.clean_text(str(d.get('codigo', ''))), 0)

        pdf.set_font("Arial", "B", 10)
        pdf.cell(30, 8, "UBICACION:", 0)
        pdf.set_font("Arial", "", 10)
        pdf.cell(50, 8, pdf.clean_text(d['gabinete']), 0, 1)

        pdf.set_font("Arial", "B", 10)
        pdf.cell(30, 8, "STOCK ACTUAL:", 0)
        pdf.set_font("Arial", "", 10)
        pdf.cell(50, 8, pdf.clean_text(str(d['cantidad'])), 0, 1)
        pdf.ln(5)

        pdf.set_font("Arial", "B", 10)
        pdf.cell(0, 8, "DESCRIPCION:", 0, 1)
        pdf.set_font("Arial", "", 10)
        pdf.multi_cell(0, 5, pdf.clean_text(d['descripcion']))
        pdf.ln(10)

        pdf.set_font("Arial", "B", 10)
        pdf.cell(0, 8, "HISTORIAL DE MOVIMIENTOS:", 0, 1)
        pdf.set_font("Arial", size=8)

        if d.get('historial', []):
            for line in d['historial']:
                pdf.cell(0, 6, pdf.clean_text(line), 1, 1)
        else:
            pdf.cell(0, 6, "Sin movimientos registrados", 1, 1)

        pdf.output(ruta)
        if messagebox.askyesno("Exito", "PDF generado.\n¿Abrir?"): os.startfile(ruta)


if __name__ == "__main__":
    root = tk.Tk();
    app = SistemaInventario(root);
    root.mainloop()